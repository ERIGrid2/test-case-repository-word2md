from typing import List
import math
import os
from datetime import date
from difflib import SequenceMatcher
from docx import Document
from docx.document import Document as Doc
import markdown
from lxml import etree
import chevron

from word2md.markdown_document import (
    MarkdownDocument, 
    MarkdownParagraph,
    MarkdownGraphic,
    MarkdownEquation,
    MarkdownTable,
    MarkdownTableRow,
    MarkdownTableCell
)

class Word2MDConverter:
    CONVERTER_TYPE = 'Test Case'

    def __init__(self, document, no_emf=False):
        if isinstance(document, Doc):
            self.document = document
        else:
            self.document = Document(document)
        self.no_emf = no_emf
        self.is_extension = False

        self.mml_transform = etree.XSLT(etree.parse(os.path.join(os.path.dirname(__file__), 'xsl', 'omml2mml_v2.xsl')))
        self.remove_namespaces = etree.XSLT(etree.parse(os.path.join(os.path.dirname(__file__), 'xsl', 'remove_namespaces.xsl')))

    def convert(self) -> List[MarkdownDocument]:
        '''
        Converts the Word document into one or more outputs.
        Each ouput has the format
        {
            'path': Relative path where the output should be stored,
            'markdown': Markdown content as string
            'attachments': List with attachments (figures) 
            [ 
                {
                    'type': Type of attachment (default 'figure'),
                    'name': Name of attachment file,
                    'data': Data of attachment
                }            
            ]
        }   
        '''
        to_md_docs = self.internal_convert()

        return to_md_docs


    def internal_convert(self) -> List[MarkdownDocument]:
        '''
        Returns a dictionary with the following:
        {
            'content': The content that should be converted to Markdown,
            'path': Path where the output should be saved
        }
        '''
        raise RuntimeError(f'{self.__class__.__name__} must implement the method {self.internal_convert.__name__}.')
    
    
    def parse_table(self, table) -> MarkdownTable:
        raw_table = MarkdownTable()
        for r, row in enumerate(table.rows):
            raw_row = MarkdownTableRow()
            for c, cell in enumerate(self.get_raw_cells(row)):
                raw_cell = MarkdownTableCell(is_heading=self.is_heading(cell['this'], r, c), paragraphs=self.get_cell_contents(cell['this']), colspan=cell['colspan'])
                raw_row.cells.append(raw_cell)
            raw_table.rows.append(raw_row)
        return raw_table
    
    def get_raw_cells(self, row):
        cells = []
        row_cells = row.cells
        for c in row_cells:
            if len(cells) == 0 or c != cells[-1]['this']:
                cell = {'this': c, 'colspan': row_cells.count(c)}
                cells.append(cell)
        return cells

    def is_heading(self, cell, row_nr, col_nr):
        raise RuntimeError(f'{self.__class__.__name__} must implement the method {self.is_heading.__name__}.')
   
    def get_paragraph_text(self, paragraph):
        prefix = ''
        if self.is_bullet_list(paragraph):
            level = self.get_numbering_level(paragraph)
            prefix = '    '.join(['' for _ in range(level)]) + '- '
        elif self.is_numbered_list(paragraph):
            level = self.get_numbering_level(paragraph)
            prefix = '    '.join(['' for _ in range(level)]) + '1. '
        return prefix + paragraph.text

    def is_bullet_list(self, paragraph):
        return self.get_numbering_format(paragraph) == 'bullet'

    def is_numbered_list(self, paragraph):
        fmt = self.get_numbering_format(paragraph)
        return fmt is not None and fmt != 'bullet'

    def get_numbering_level(self, paragraph):
        lvl = self.get_numbering_lvl(paragraph)
        if lvl is not None:
            try:
                lvl_indent = lvl.find('w:pPr/w:ind', namespaces=lvl.nsmap)
                if lvl_indent is not None:
                    level = math.floor(lvl_indent.left / 500000) + 1
                else:
                    level = int(self.get_value_of_attribute(lvl, 'ilvl')) + 1

                return level
            except:
                pass
        return 0

    def get_numbering_lvl(self, paragraph):
        document_part = paragraph.part
        namespaces = paragraph._element.nsmap
        w_namespace = namespaces['w']
        p_numbering = paragraph._element.find('*/w:numPr', namespaces=namespaces)
        if p_numbering is not None:
            att_val = '{' + w_namespace + '}val'
            ilvl = p_numbering.find('w:ilvl', namespaces=namespaces)
            numId = p_numbering.find('w:numId', namespaces=namespaces)
            if ilvl is not None and numId is not None:
                abstract_num_id = document_part.numbering_part.element.find(
                    'w:num[@w:numId="' + self.get_attr_val(numId) + '"]/w:abstractNumId', 
                    namespaces=namespaces)
                if abstract_num_id is not None:
                    xpath_str = ('w:abstractNum[@w:abstractNumId="' + self.get_attr_val(abstract_num_id) + '"]' + 
                        '/w:lvl[@w:ilvl="' + self.get_attr_val(ilvl) + '"]')
                    num_level = document_part.numbering_part.element.find(xpath_str, namespaces=namespaces)
                    return num_level
        return None

    def get_attr_val(self, element):
        return self.get_value_of_attribute(element, 'val')

    def get_value_of_attribute(self, element, attribute_name):
        namespaces = element.nsmap
        w_namespace = namespaces['w']
        attribute = '{' + w_namespace + '}' + attribute_name
        return element.get(attribute)

    def get_numbering_format(self, paragraph):
        num_level = self.get_numbering_lvl(paragraph)
        if num_level is not None:
            xpath_str = 'w:numFmt'
            num_fmt = num_level.find(xpath_str, namespaces=num_level.nsmap)
            if num_fmt is not None:
                return self.get_attr_val(num_fmt)
        return None    

    def get_cell_contents(self, cell, lineseparator='\n') -> List[MarkdownParagraph]:
        contents = self.get_content_from_paragraphs(cell.paragraphs, lineseparator=lineseparator)

        for t in cell.tables:
            for c in t._cells:
                contents.extend(self.get_cell_contents(c, lineseparator=lineseparator))

        return contents
    
    def get_content_from_paragraphs(self, paragraphs, lineseparator='\n') -> List[MarkdownParagraph]:
        contents = []
        current_list_paragraph = []

        for paragraph in paragraphs:
            if self.is_numbered_list(paragraph) or self.is_bullet_list(paragraph):
                current_list_paragraph.append(self.get_paragraph_text(paragraph))
                continue
            
            if current_list_paragraph:
                p = MarkdownParagraph(text=lineseparator.join(current_list_paragraph))
                current_list_paragraph = []
                contents.append(p)

            p = MarkdownParagraph()
            p.text = self.get_paragraph_text(paragraph)            
            p.graphics.extend(self.get_inline_graphics(paragraph, self.document))            
            p.equations.extend(self.get_inline_equations(paragraph))

            contents.append(p)

        if current_list_paragraph:
            contents.append(MarkdownParagraph(text=lineseparator.join(current_list_paragraph)))

        return contents

    def get_inline_graphics(self, word_part, document) -> List[MarkdownGraphic]:
        try:
            element = word_part._element
        except:
            return []

        drawing_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"

        image_parts = []
        graphics = []
        for drawing in element.findall('*//w:drawing', namespaces=element.nsmap):
            namespaces = element.nsmap
            if drawing_ns not in namespaces.values():
                namespaces['a'] = drawing_ns
            blip = drawing.find('*//a:blip[@r:embed]', namespaces=namespaces)
            if blip is not None:
                graphic_id = blip.embed
                image_part = document.part.related_parts[graphic_id]         
                image_parts.append(image_part)
        
        for imagedata in element.findall('*//v:imagedata', namespaces=document._element.nsmap):
            image_id = imagedata.get('{{{0}}}id'.format(imagedata.nsmap['r']))
            image_part = document.part.related_parts[image_id]                
            image_parts.append(image_part)

        for image_part in image_parts:
            image_path = image_name = os.path.basename(image_part.partname)
            if self.no_emf:
                if '.' in image_name and image_name.endswith('.emf'):
                    image_name = '.'.join(image_name.split('.')[:-1]) + '.png'
            graphic = MarkdownGraphic(name=image_name, src=image_path, data=image_part._blob)
            graphics.append(graphic)
            
        return graphics

    def get_inline_equations(self, word_part) -> List[MarkdownEquation]:
        try:
            element = word_part._element
        except:
            return []
        
        equations = []
        for equation in element.findall('*//m:oMath', namespaces=element.nsmap):
            eq = {}
            mml = self.mml_transform(equation)
            mml = self.remove_namespaces(mml.getroot())
            mml = etree.tostring(mml, encoding='unicode')
            eq = MarkdownEquation(mml=mml)
            equations.append(eq)

        return equations    

    def is_bold(self, paragraph):
        if paragraph.text:
            for r in paragraph.runs:
                if r.text and not r.bold:
                    return False
            return True 
        return False       
    